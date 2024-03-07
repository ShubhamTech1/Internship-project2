



                                            # EDA CODE
                                            
                                            
#========================================================================================================================================

# step 1st:- drop unwanted columns


#pip install pandas
#pip install sqlalchemy
#pip install matplotlib
#pip install scipy
#pip install scikit-learn
#pip install seaborn
#pip install statsmodels


import pandas as pd                                        ## data manipulation
import numpy as np                                         ## numerical calculation 
from sklearn.preprocessing import StandardScaler
from sklearn.preprocessing import PowerTransformer
from sklearn.impute import SimpleImputer
import matplotlib.pyplot as plt                            ## data visualization
from sqlalchemy import create_engine                       ## connect to SQL database
from feature_engine.outliers import Winsorizer
import seaborn as sns
import statsmodels.api as sm
import scipy.stats as stats
import pylab
from scipy import stats
from scipy.stats import skew, kurtosis, mode






df1 = pd.read_csv(r"D:\360DigiTMG\PROJECTS\2nd project\Project Materials\Dataset\kit.csv")

# MySQL Database connection
# Creating engine which connect to MySQL
user = 'grant3' # user name
pw = 'grant3' # password
db = 'kititems' # database

# creating engine to connect database
engine = create_engine(f"mysql+pymysql://{user}:{pw}@localhost/{db}")

# dumping data into database 
df1.to_sql('kit', con = engine, if_exists = 'replace', chunksize = 1000, index = False)

# loading data from database
sql = 'select * from kit'
df = pd.read_sql_query(sql, con = engine)








df = pd.read_csv(r"D:\360DigiTMG\PROJECTS\2nd project\Project Materials\Dataset\kit.csv")
#Deleting Unnecessory columns
Data = df.drop(columns={'Customer Code', 'Customer Name','OEM', 'Item Description','Product type','Total','Item Code'})


# Set the 'KIT ITEM' column as the index
Data.set_index('KIT ITEM', inplace=True)

# Transpose the DataFrame to swap rows and columns
Data = Data.transpose()

Data.columns

Data.shape

Data.info()

# Null Value Checking
Data.isnull().sum()

Desc = Data.describe()

# Identify duplicate column names
duplicate_columns = Data.columns[Data.columns.duplicated()].tolist() 
duplicate_columns
    
#removing the duplicate column 
Data = Data.loc[:, ~Data.columns.duplicated()] 

# FIRST MOMENT BUSINESS DECISION
mean = Data.mean() 
median = Data.median() 
mode = Data.mode()

# SECOND MOMENT BUSINESS DECISION
variance = Data.var()
satndard_deviation = Data.std()

# THIRD MOMENT BUSINESS DECISION
Skewness= Data.skew()

# FOURTH MOMENT BUSINESS DECISION
Kurtosis= Data.kurt()

#Removing near zero variance features

from sklearn.feature_selection import VarianceThreshold

var_thres=VarianceThreshold(threshold=0)
var_thres.fit(Data)

var_thres.get_support()
Data.columns[var_thres.get_support()]

constant_columns = [column for column in Data.columns
                    if column not in Data.columns[var_thres.get_support()]]

print(len(constant_columns))

for feature in constant_columns:
     print(feature)
     
    
Data = Data.drop(constant_columns,axis = 1)

Data.shape

#data=data.drop(data.columns[data.isnull().mean()<=0.20], axis=1)
#data = data.dropna(axis=1, thresh= 0.80)

## fill nan with zero ## 
Data = Data.fillna(0) 

# Define threshold for proportion of zeros
zero_threshold = 0.20 # Columns with >80% zeros will be removed

# Calculate proportion of zeros in each column
zero_proportion = (Data == 0).mean()

# Filter columns based on zero proportion threshold
high_zero_cols = zero_proportion[zero_proportion > zero_threshold].index

# Remove columns with high proportion of zeros
Data1 = Data.drop(columns = high_zero_cols)

# Replacing 0 with nan values
Data2 = Data1.replace(0, np.nan)

#Saving data after deleting column
Data2.to_csv('KIT_ITEM2.CSV')

import os
os.getcwd()


#========================================================================================================================================



import pandas as pd

#  dataset is stored in a DataFrame called df
# Drop columns with at least one null value in the last 6 rows

cols_with_null = Data2.iloc[-6:].isnull().any()

new_df = Data2.loc[:, ~cols_with_null]

new_df.to_csv('KIT_ITEM1.CSV')

import os
os.getcwd()


#========================================================================================================================================




# step 2nd :- visualize data after dropping unwanted columns


Datas = pd.read_csv(r"C:\Users\Asus\KIT_ITEM1.CSV") 

# Select the 'Total' column
values = Datas['Total']

# Select the 'Dates' column
dates = Datas['Dates']

# Plot the time series
plt.figure(figsize=(10, 6))
#plt.plot(x_axis = dates,y_axis = values, marker='o', color='blue', linestyle='-')
plt.plot(dates, values, color='blue', linestyle='-')
plt.title('Time Series Plot')
plt.xlabel('Date')
plt.ylabel('Values')

# Show plot
plt.show()







# step 3rd :- missing values impute using KNN interpolation




import pandas as pd
from sklearn.impute import KNNImputer

# Read the CSV file into a pandas DataFrame
Data3 = pd.read_csv(r"C:\Users\Asus\KIT_ITEM1.CSV") 

# Convert the 'DATES' column to datetime objects with the specified format
Data3['Dates'] = pd.to_datetime(Data3['Dates'], format='%m-%d-%Y %H.%M') 

Data3.dtypes 
# Drop non-numeric columns (if any) 
data_numeric = Data3.select_dtypes(include=['float64', 'int64'])   

# Initialize KNNImputer
imputer = KNNImputer(n_neighbors = 5)

# Perform imputation on numeric data only
imputed_data = imputer.fit_transform(data_numeric)

# Convert Imputed Data to DataFrame
imputed_df = pd.DataFrame(imputed_data, columns=data_numeric.columns)

# Round the imputed values to the nearest whole number
imputed_df = imputed_df.round().astype(int)

# Insert 'Dates' column at the beginning of the DataFrame
imputed_df.insert(0, 'Dates', Data3['Dates'])

#Saving data after deleting column
imputed_df.to_csv('knnimputed.CSV', index = False) 

import os
os.getcwd()




#========================================================================================================================================

# step 4th :- outlier treatment
import matplotlib.pyplot as plt

Clean = pd.read_csv(r"C:\Users\Asus\knnimputed.CSV")  
Clean.plot(kind = 'box', subplots = True, sharey = False, figsize = (15, 8)) 
# outliers present in dataset


from feature_engine.outliers import Winsorizer

# Segregating Numeric features
numeric_features = Clean.select_dtypes(exclude=['object']).columns.tolist()

# Winsorization for outlier treatment on numerical features only
winsor = Winsorizer(
    capping_method ='iqr',  # choose IQR rule boundaries or gaussian for mean and std
    tail ='both',  # cap left, right or both tails
    fold =1.5,
    variables = numeric_features)  # Apply only to numerical columns

clean_data = winsor.fit_transform(Clean)  # Fit and transform numerical features
clean_data.to_csv("clean_data.csv")


                        ## Visualize data after KNNImputer and outlier treatment:-
                         
import matplotlib.pyplot as plt

VS = pd.read_csv(r"C:\Users\Asus\clean_data.csv")  
VS.plot(kind = 'box', subplots = True, sharey = False, figsize = (15, 8)) 

# Select the 'Total' column
values = VS['Total']

# Select the 'Dates' column
dates = VS['Dates']

# Plot the time series
plt.figure(figsize=(10, 6))
#plt.plot(x_axis = dates,y_axis = values, marker='o', color='blue', linestyle='-')
plt.plot(dates, values, color='blue', linestyle='-')
plt.title('Time Series Plot')
plt.xlabel('Date')
plt.ylabel('Values')

# Show plot
plt.show()








# missing values impute using Linear Interpolation

import pandas as pd 
Data3 = pd.read_csv(r"D:\360DigiTMG\PROJECTS\2nd project\Answer Sheet\Python\KIT_ITEM2.CSV") 

# Assuming Data3 is your DataFrame
Data3.interpolate(method='linear', inplace=True) 

# Convert Data3 back to a DataFrame
Data3 = pd.DataFrame(Data3)

# Alternatively, you can assign it to a new variable
interpolated_df = pd.DataFrame(Data3) 

# not doing outlier treatment and visualize data after fill missing values using Interpolation



#========================================================================================================================================


# VISUALIZE DATA IN WORD DOCUMENT BEFORE PREPROCESSING

import pandas as pd
from sklearn.impute import KNNImputer
import numpy as np
from sqlalchemy import create_engine
from urllib.parse import quote 
from getpass import getpass
from statsmodels.tsa.stattools import adfuller, kpss
from statsmodels.tsa.seasonal import seasonal_decompose
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document

# Read the CSV file into a pandas DataFrame
Data3 = pd.read_csv(r"D:\360DigiTMG\PROJECTS\2nd project\Answer Sheet\Python\KIT_ITEM2.CSV") 

# Convert the 'DATES' column to datetime objects with the specified format
Data3['Dates'] = pd.to_datetime(Data3['Dates'], format='%m-%d-%Y %H.%M') 

Data3.dtypes 
# Drop non-numeric columns (if any) 
data_numeric = Data3.select_dtypes(include=['float64', 'int64'])   

# Initialize KNNImputer
imputer = KNNImputer(n_neighbors = 5)

# Perform imputation on numeric data only
imputed_data = imputer.fit_transform(data_numeric)

# Convert Imputed Data to DataFrame
imputed_df = pd.DataFrame(imputed_data, columns=data_numeric.columns)

# Round the imputed values to the nearest whole number
data = imputed_df.round().astype(int)

# Insert 'Dates' column at the beginning of the DataFrame
data.insert(0, 'Dates', Data3['Dates'])


# Create a new Word document
doc = Document()

# Iterate over each column (time series) in the DataFrame
for column in data.columns:
    doc.add_heading(f"Stationary tests for {column}", level=1)

    # ADF Test
    adf_result = adfuller(data[column])
    doc.add_paragraph(f"ADF Test - p-value: {adf_result[1]}")
    if adf_result[1] <= 0.05:
        doc.add_paragraph("ADF Test: Series is stationary")
    else:
        doc.add_paragraph("ADF Test: Series is not stationary")

    # KPSS Test
    kpss_result = kpss(data[column])
    doc.add_paragraph(f"KPSS Test - p-value: {kpss_result[1]}")
    if kpss_result[1] >= 0.05:
        doc.add_paragraph("KPSS Test: Series is stationary")
    else:
        doc.add_paragraph("KPSS Test: Series is not stationary")

    # Seasonality Trend Decomposition Plot
    decomposition = seasonal_decompose(data[column], model='additive', period=12)  # Change period accordingly
    plt.figure(figsize=(10, 8))  # Adjust figure size for better readability
    plt.subplot(411)
    plt.plot(decomposition.observed)
    plt.ylabel('Original')
    plt.xticks(rotation=45)  # Rotate x-axis labels for better visibility
    plt.tight_layout()  # Adjust layout

    plt.subplot(412)
    plt.plot(decomposition.trend)
    plt.ylabel('Trend')
    plt.xticks(rotation=45)
    plt.tight_layout()

    plt.subplot(413)
    plt.plot(decomposition.seasonal)
    plt.ylabel('Seasonal')
    plt.xticks(rotation=45)
    plt.tight_layout()

    plt.subplot(414)
    plt.plot(decomposition.resid)
    plt.ylabel('Residual')
    plt.xticks(rotation=45)
    plt.tight_layout()

    # Adjust spacing between subplots
    plt.subplots_adjust(hspace=1.5)

    # Create a BytesIO object to temporarily store the plot
    tmp_img = BytesIO()
    plt.savefig(tmp_img, format='png')
    tmp_img.seek(0)

    # Add the plot to the Word document
    doc.add_picture(tmp_img, width=Inches(5))
    plt.close()

    # Add a page break for each column
    doc.add_page_break()

# Save the Word document
doc.save("stationarity_tests_output.docx")




